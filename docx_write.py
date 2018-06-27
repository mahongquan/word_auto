# -*- coding: utf-8 -*-
import pickle
from docx import Document
import logging
addSerial=0
def addStart(start):
    global addSerial
    addSerial+=1
    if addSerial % 2==0:
        start+=1
    return start
def change(start,p0):
    i=0
    for r in p0.runs:
        print(i,r.text)
        i+=1
    i+=1
    n=len(p0.runs)
    p0.runs[5].text="No  0018%04d" %start
    i=6
    while(i<n):
        p0.runs[i].text=""
        i+=1
    start=addStart(start)
    return start
def genPack0(start,fn):
    document = Document(fn)
    print(dir(document))
    input("aaaaaa")
    for p0 in document.paragraphs:
        if len(p0.runs)>0:
            if p0.runs[0].text=="北京科技大学":
               start=change(start,p0)
    document.save("新建.docx")
    pickle.dump(start, open("start.pickle","wb"), protocol=None)    
def genPack(start,fn):
    document = Document(fn)
    print(dir(document))
    input("aaaaaa")
    for p0 in document.paragraphs:
        if len(p0.runs)>0:
            if p0.runs[0].text=="北京科技大学":
               start=change(start,p0)
    document.save("新建.docx")
    pickle.dump(start, open("start.pickle","wb"), protocol=None)
if __name__=="__main__":
    print(genPack("4111533499"))